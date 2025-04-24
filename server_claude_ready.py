from fastmcp import FastMCP
import time
import logging
from pymilvus import Collection, connections, utility, FieldSchema, CollectionSchema, DataType

# Setup FastMCP
mcp = FastMCP()

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("milvus_tools")

# Connect to Milvus
for i in range(10):
    try:
        connections.connect(alias="default", host="localhost", port="19530")
        logger.info("Connected to Milvus at standalone:19530")
        break
    except Exception as e:
        logger.warning(f"Waiting for Milvus... ({i+1}/10): {e}")
        time.sleep(2)
else:
    raise ConnectionError("Failed to connect to Milvus after multiple attempts.")

# Schema builder
def get_default_schema(dim: int = 1536):
    fields = [
        FieldSchema(name="id", dtype=DataType.INT64, is_primary=True, auto_id=True),
        FieldSchema(name="embedding", dtype=DataType.FLOAT_VECTOR, dim=dim)
    ]
    return CollectionSchema(fields, description="Default schema")


@mcp.tool(name="Create basic Milvus collection")
def create_collection(name: str, dim: int = 1536) -> str:
    """Create a new Milvus collection."""
    if name in utility.list_collections():
        return f"Collection '{name}' already exists."
    schema = get_default_schema(dim)
    Collection(name, schema)
    return f"Collection '{name}' created."


@mcp.tool(name="Insert vector data into a collection")
def insert_data(collection_name: str, vectors: list) -> str:
    """Insert vectors into a Milvus collection."""
    collection = Collection(collection_name)
    entities = [vectors]
    collection.insert(entities)
    return f"Inserted {len(vectors)} vectors into '{collection_name}'."


@mcp.tool(name="Create search index for a collection")
def create_index(collection_name: str) -> str:
    """Create an index on a Milvus collection."""
    collection = Collection(collection_name)
    index_params = {
        "index_type": "IVF_FLAT",
        "metric_type": "L2",
        "params": {"nlist": 128}
    }
    collection.create_index(field_name="embedding", index_params=index_params)
    return f"Index created on '{collection_name}'."


@mcp.tool(name="Load collection into memory")
def load_collection(collection_name: str) -> str:
    """Load a Milvus collection into memory."""
    collection = Collection(collection_name)
    collection.load()
    return f"Collection '{collection_name}' loaded."


@mcp.tool(name="Search vectors in a collection")
def search(collection_name: str, query_vector: list, top_k: int = 5) -> list:
    """Search Milvus collection with a query vector."""
    collection = Collection(collection_name)
    collection.load()  # load is safe to call even if already loaded

    search_params = {"metric_type": "L2", "params": {"nprobe": 10}}
    results = collection.search(
        data=[query_vector],
        anns_field="embedding",
        param=search_params,
        limit=top_k
    )

    return [{"id": hit.id, "score": hit.distance} for hit in results[0]]

@mcp.tool(name="Upload document using ASCII embedding")
def upload_document_ascii(collection_name: str, text: str, doc_id: str = None) -> str:
    """Upload a long document by splitting, embedding, and storing chunks with metadata."""
    import uuid

    # Dummy embedder: char-to-ASCII embedding
    def embed(text):
        vec = [float(ord(c)) / 255.0 for c in text]
        if len(vec) > 1536:
            return vec[:1536]
        return vec + [0.0] * (1536 - len(vec))

    # Simple splitter: every 500 characters
    def split_into_chunks(text, chunk_size=500):
        return [text[i:i + chunk_size] for i in range(0, len(text), chunk_size)]

    chunks = split_into_chunks(text)
    vectors = [embed(chunk) for chunk in chunks]
    texts = chunks
    doc_ids = [doc_id or str(uuid.uuid4())] * len(chunks)

    collection = Collection(collection_name)

    # Add a 'doc_id' field to the schema manually if needed
    # Skipped here – assumes collection was created with text and doc_id fields

    collection.insert([vectors, texts, doc_ids])
    create_index(collection_name)

    return f"Uploaded document with {len(chunks)} chunks to '{collection_name}'."


@mcp.tool(name="Create basic Milvus collection")
def create_collection_with_doc_id(name: str, dim: int = 1536) -> str:
    """Create a Milvus collection with 'text' and 'doc_id' metadata fields."""
    if name in utility.list_collections():
        return f"Collection '{name}' already exists."

    fields = [
        FieldSchema(name="id", dtype=DataType.INT64, is_primary=True, auto_id=True),
        FieldSchema(name="embedding", dtype=DataType.FLOAT_VECTOR, dim=dim),
        FieldSchema(name="text", dtype=DataType.VARCHAR, max_length=1024),
        FieldSchema(name="doc_id", dtype=DataType.VARCHAR, max_length=128),
    ]
    schema = CollectionSchema(fields, description="Embedding + metadata collection")
    Collection(name, schema)
    return f"Collection '{name}' with metadata fields created."


@mcp.tool(name="Extract text from a file")
def extract_text_from_file(file_path: str) -> str:
    """Extracts plain text from a file (PDF, DOCX, or TXT)."""
    import os
    from pathlib import Path

    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    elif ext == ".docx":
        import docx
        doc = docx.Document(file_path)
        return "\n".join(p.text for p in doc.paragraphs)

    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    else:
        return f"Unsupported file type: {ext}"


@mcp.tool(name="Upload document using OpenAI embedding")
def upload_document_openai(collection_name: str, text: str, doc_id: str = None) -> str:
    """Split, embed with OpenAI, and store a document in Milvus with metadata."""
    import uuid
    import openai
    import os

    openai.api_key = os.getenv("OPENAI_API_KEY")

    def split_into_chunks(text, chunk_size=500):
        return [text[i:i + chunk_size] for i in range(0, len(text), chunk_size)]

    from openai import OpenAI
    client = OpenAI()


    def embed_texts(text_list):
        response = client.embeddings.create(
            model="text-embedding-ada-002",
            input=text_list
        )

        vectors = [d.embedding for d in response.data]

    chunks = split_into_chunks(text)
    vectors = embed_texts(chunks)
    texts = chunks
    doc_ids = [doc_id or str(uuid.uuid4())] * len(chunks)

    collection = Collection(collection_name)
    collection.insert([vectors, texts, doc_ids])
    create_index(collection_name)

    return f"Uploaded document with {len(chunks)} embedded chunks to '{collection_name}'."


@mcp.tool(name="Ingest and embed file into Milvus")
def ingest_file_to_milvus(file_path: str, collection_name: str, doc_id: str = None) -> str:
    """One-click tool: extract text from a file, embed, and store in Milvus."""
    import uuid
    import openai
    import os
    from pathlib import Path

    openai.api_key = os.getenv("OPENAI_API_KEY")

    # 1. Extract text
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
    elif ext == ".docx":
        import docx
        doc = docx.Document(file_path)
        text = "\n".join(p.text for p in doc.paragraphs)
    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
    else:
        return f"Unsupported file type: {ext}"

    # 2. Chunk
    def split_into_chunks(text, chunk_size=500):
        return [text[i:i + chunk_size] for i in range(0, len(text), chunk_size)]

    chunks = split_into_chunks(text)
    if not chunks:
        return "No text found in file."

    # 3. Embed
    from openai import OpenAI
    client = OpenAI()

    def embed_texts(text_list):
        response = client.embeddings.create(
            model="text-embedding-ada-002",
            input=text_list
        )
        return [d.embedding for d in response.data]

    vectors = embed_texts(chunks)
    doc_ids = [doc_id or str(uuid.uuid4())] * len(chunks)

    # 4. Insert
    collection = Collection(collection_name)
    collection.insert([vectors, chunks, doc_ids])
    create_index(collection_name)

    return f"Ingested {len(chunks)} chunks from '{file_path}' into '{collection_name}'."


@mcp.tool(name="Ask a question about documents")
def query_document_summary(collection_name: str, query: str, top_k: int = 5) -> str:
    """Search Milvus for similar chunks and return a concatenated summary string."""
    from openai import OpenAI
    client = OpenAI()

    # Embed the query using OpenAI v1.0+
    response = client.embeddings.create(
        model="text-embedding-ada-002",
        input=[query]
    )
    query_vector = response.data[0].embedding

    # Search Milvus
    collection = Collection(collection_name)
    collection.load()
    results = collection.search(
        data=[query_vector],
        anns_field="embedding",
        param={"metric_type": "L2", "params": {"nprobe": 10}},
        limit=top_k,
        output_fields=["text"]
    )

    # Return the top matching chunks
    summary = "\n---\n".join(hit.entity.get("text") for hit in results[0])
    return f"Top {top_k} matching excerpts:\n{summary}"


@mcp.tool(name="List all Milvus collections")
def list_collections() -> list:
    """List all available Milvus collections."""
    return utility.list_collections()


@mcp.tool(name="Check Milvus connection health")
def health_check() -> str:
    """Check if Milvus is connected and healthy."""
    try:
        utility.list_collections()
        return "Milvus is connected and ready!"
    except Exception as e:
        return f"Milvus connection error: {e}"


if __name__ == "__main__":
    mcp.run()
